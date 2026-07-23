from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for level in [1, 2, 3]:
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    if level == 1:
        h.font.size = Pt(22)
        h.font.bold = True
    elif level == 2:
        h.font.size = Pt(15)
        h.font.bold = True
    elif level == 3:
        h.font.size = Pt(11)
        h.font.bold = True
        h.font.color.rgb = RGBColor(0x40, 0x40, 0x60)

def add_page_break():
    doc.add_page_break()

def add_goal_logic(goal, logic):
    p = doc.add_paragraph()
    r = p.add_run('Goal: ')
    r.bold = True
    r.font.size = Pt(10.5)
    r = p.add_run(goal)
    r.font.size = Pt(10.5)
    p2 = doc.add_paragraph()
    r = p2.add_run('Logic: ')
    r.bold = True
    r.font.size = Pt(10.5)
    r = p2.add_run(logic)
    r.font.size = Pt(10.5)

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1A1A2E')
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_shading(cell, 'F0F0F5')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def add_checklist(doc, items):
    for item in items:
        p = doc.add_paragraph()
        r = p.add_run(f'\u2610  {item}')
        r.font.size = Pt(10)

def add_field(label, width_cm=8):
    p = doc.add_paragraph()
    r = p.add_run(f'{label}: ')
    r.bold = True
    r.font.size = Pt(10)
    r = p.add_run('_' * int(width_cm * 3))
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

def add_mono_block(doc, text, size=8.5):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(size)

def add_small_heading(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

# ═══════════════════════════════════════
# SECTION 1 — Cover + Overview
# ═══════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('CONCERTO EVENT PROMPT BOOK')
r.font.size = Pt(18)
r.bold = True
r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(6)
r = p.add_run('─' * 50)
r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_p.paragraph_format.space_after = Pt(8)
meta_items = [
    ('Project', '[Composer — Concerto Name]'),
    ('Soloist', '[Name]'),
    ('Conductor', '[Name]'),
    ('Orchestra', '[Name]'),
    ('Venue', '[Venue]'),
    ('Dates', '[Rehearsal — Recording]'),
    ('Version', '[Draft / Final]'),
]
for i, (label, val) in enumerate(meta_items):
    r = meta_p.add_run(f'{label}: ')
    r.bold = True
    r.font.size = Pt(9.5)
    r = meta_p.add_run(val)
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    if i < len(meta_items) - 1:
        r = meta_p.add_run('    ')
        r.font.size = Pt(9.5)

doc.add_heading('1. Event Overview — How to Use This Book', level=2)

add_goal_logic(
    'Explain the concerto project, the recording format, and how this book maps the entire production pipeline.',
    'This is a production document, not a musical score. It tracks logistics from first rehearsal through final patched master. A new producer, engineer, or orchestra manager should open this page and immediately understand the project scope, the timeline, and where to find each piece of operational detail.'
)

add_field('Concerto / Work(s) Being Recorded')
add_field('Duration (planned)')
add_field('Recording Format (live / takes / hybrid)')
add_field('Distribution / Use Case')

make_table(doc, ['#', 'Section', 'When You Need It'], [
    ['1', 'Event Overview', 'Onboarding anyone new to the project'],
    ['2', 'People Involved', 'Contacting anyone, understanding chain of command'],
    ['3', 'Scheduling', 'Building or checking the timeline at any phase'],
    ['4', 'Orchestra Setup', 'Load-in day — stage plot, approvals, physical setup'],
    ['5', 'Rehearsal Week', 'Running rehearsals day by day'],
    ['6', 'In Between Weekend', 'Gap-day logistics — chairs, doors, audience prep'],
    ['7', 'Recording Week', 'Running recording sessions day by day'],
    ['8', 'Video Adjustment', 'Pre-recording camera/lighting/positioning setup'],
    ['9', 'Recording', 'Session procedure — takes, roll numbers, comms'],
    ['10', 'Patching', 'Post-production — comping takes, syncing, deliverables'],
    ['11', 'Unexpected Changes', 'Contingency plans for common failure modes'],
], col_widths=[1, 6, 9.5])

# ═══════════════════════════════════════
# SECTION 2
# ═══════════════════════════════════════
add_page_break()
doc.add_heading('2. People Involved', level=2)

add_goal_logic(
    'List every person on the production with their role, contact, and who they report to.',
    'On a recording session, time is money — every musician on the clock. Wasting ten minutes finding the right person to approve a chair move costs real dollars. One page with every contact means the producer solves problems without leaving their seat.'
)

make_table(doc, ['Role', 'Name', 'Phone', 'Email', 'Reports To'], [
    ['Artistic Director / Producer', '', '', '', ''],
    ['Conductor', '', '', '', ''],
    ['Soloist', '', '', '', ''],
    ['Orchestra Manager', '', '', '', ''],
    ['Personnel Manager', '', '', '', ''],
    ['Librarian', '', '', '', ''],
    ['Recording Producer', '', '', '', ''],
    ['Recording Engineer (Audio)', '', '', '', ''],
    ['Recording Engineer (Video)', '', '', '', ''],
    ['Video Director', '', '', '', ''],
    ['Lighting Designer', '', '', '', ''],
    ['Stage Manager', '', '', '', ''],
    ['Stage Crew Lead', '', '', '', ''],
    ['Editor / Patching Engineer', '', '', '', ''],
    ['Venue Coordinator', '', '', '', ''],
    ['Union Steward (if applicable)', '', '', '', ''],
    ['PR / Communications', '', '', '', ''],
], col_widths=[4.5, 2.5, 2.5, 2.5, 2.5])

add_small_heading('Chain of Command (on session floor)')
add_mono_block(doc, '''Producer \u2192 Conductor \u2192 Orchestra Manager \u2192 Personnel Manager \u2192 Musicians
                  \u2192 Recording Engineer \u2192 Assistant Engineer
                  \u2192 Stage Manager \u2192 Stage Crew
Illustration: In session, any non-musical issue routes through the Orchestra Manager,
not the Conductor. The Conductor focuses on the performance. The Producer runs the room.''', size=9)

add_small_heading('Emergency Contacts')
make_table(doc, ['Type', 'Name', 'Phone'], [
    ['Venue Emergency', '', ''],
    ['Medical / First Aid', '', ''],
    ['Fire / Security', '', ''],
    ['Union Hotline', '', ''],
    ['Transport / Parking', '', ''],
], col_widths=[5.5, 5.5, 5.5])

# ═══════════════════════════════════════
# SECTION 3
# ═══════════════════════════════════════
add_page_break()
doc.add_heading('3. Scheduling', level=2)

add_goal_logic(
    'Show the entire timeline in one view: rehearsal week \u2192 gap weekend \u2192 recording week \u2192 post-production.',
    'A concerto recording has three distinct phases with different teams active. The schedule must show who needs to be where and when, including the "gap" days that look empty but are critical for setup changes, rest, and contingency.'
)

add_small_heading('Master Timeline')
make_table(doc, ['Phase', 'Date(s)', 'Key Activities', 'Teams Involved'], [
    ['Rehearsal Week', '', 'Sectionals, tutti rehearsals, soloist run-throughs', 'Conductor, orchestra, soloist, librarian'],
    ['In Between Weekend', '', 'Chair reconfiguration, door/audience prep, equipment check', 'Stage crew, audio/video techs'],
    ['Recording Week', '', 'Session blocks, takes, video capture', 'Full production + orchestra + soloist'],
    ['Post-Production', '', 'Patching, mixing, video sync, QC', 'Editor, mixing engineer, producer'],
    ['Delivery', '', 'Final master approved, deliverables sent', 'Producer, label/client'],
], col_widths=[3.5, 2.5, 5.5, 4.5])

add_small_heading('Rehearsal Week — Day by Day')
make_table(doc, ['Day', 'Time', 'Activity', 'Location', 'Notes'], [
    ['Mon', '', 'Setup / stage plot / sound check', '', ''],
    ['Mon', '', 'Sectional — strings', '', ''],
    ['Tue', '', 'Tutti rehearsal (no soloist)', '', ''],
    ['Tue', '', 'Orchestra soloist run (movt 1)', '', ''],
    ['Wed', '', 'Full run-through (all movements)', '', ''],
    ['Wed', '', 'Notes + retakes on problem passages', '', ''],
    ['Thu', '', 'Dress rehearsal (as live)', '', ''],
    ['Thu', '', 'Notes + final adjustments', '', ''],
    ['Fri', '', 'Make-up / overflow / rest day', '', ''],
], col_widths=[2, 2, 5, 3.5, 3.5])

add_small_heading('Recording Week — Day by Day')
make_table(doc, ['Day', 'AM Session', 'PM Session', 'Evening', 'Notes'], [
    ['Mon', 'Setup video/audio', 'Recording block 1', 'Break / review', ''],
    ['Tue', 'Recording block 2', 'Recording block 3', 'Playback review', ''],
    ['Wed', 'Recording block 4', 'Recording block 5', 'Break / review', ''],
    ['Thu', 'Pick-ups / retakes', 'Tutti patches', 'Orchestra release', ''],
    ['Fri', 'Soloist pick-ups', 'Final QC playthrough', 'Wrap', ''],
], col_widths=[1.5, 3.5, 3.5, 3.5, 3.5])

add_field('Union / Contractual Break Requirements')
add_field('Overtime Approval Chain')

add_small_heading('Schedule Attempts Log')
make_table(doc, ['Attempt #', 'Date', 'Key Change', 'Why It Changed'], [
    ['1', '', '(Initial draft)', ''],
    ['2', '', '', ''],
    ['3', '', '', ''],
    ['Final', '', '', ''],
], col_widths=[2, 2.5, 5.5, 5.5])

# ═══════════════════════════════════════
# SECTION 4
# ═══════════════════════════════════════
add_page_break()
doc.add_heading('4. Orchestra Setup', level=2)

add_goal_logic(
    'Define the stage layout, get it approved by everyone who matters, then execute setup to spec.',
    'Orchestra setup affects balance, sightlines, video framing, and musician comfort. Changing it after load-in costs hours. The map must be approved by conductor, recording engineer, video director, and orchestra manager before a single stand is placed.'
)

add_small_heading('Stage Map — Orchestra Layout')
add_mono_block(doc, '''                              [CONDUCTOR]
                                |
                       [SOLOIST]
                     (front center)
        ┌───────────────────────────────────────────────┐
        │   Vln I  │  Vln II  │  Vla    │  Vc    │  Cb  │
        │   (6-8)  │  (6-8)   │  (4-6)  │ (4-6)  │ (2-3)│
        ├──────────┼──────────┼─────────┼────────┼──────┤
        │    Fl    │   Ob     │   Cl    │   Bn   │  Hrn │
        │   (2-3)  │  (2-3)   │  (2-3)  │  (2-3) │ (3-4)│
        ├──────────┼──────────┼─────────┼────────┼──────┤
        │   Tpt    │   Tbn    │   Tba   │  Perc  │  Hp  │
        │   (2-3)  │  (2-3)   │  (1)    │  (2-3) │  (1) │
        └───────────────────────────────────────────────┘
                    [AUDIENCE / CONTROL ROOM]''', size=8)

make_table(doc, ['Measurement / Spec', 'Value', 'Approver'], [
    ['Stage width', '__________________', 'Venue'],
    ['Stage depth', '__________________', 'Venue'],
    ['Riser height (string section)', '__________________', 'Cond. + Video'],
    ['Soloist position (distance from conductor)', '__________________', 'Cond. + Soloist'],
    ['Conductor podium height', '__________________', 'Conductor'],
    ['First chair distance from conductor', '__________________', 'Conductor'],
    ['Percussion enclosure', '__________________', 'Audio + Cond.'],
], col_widths=[6, 5, 4.5])

add_small_heading('Approval Chain')
add_checklist(doc, [
    'Stage map drafted and shared (producer)',
    'Conductor approves layout',
    'Recording engineer approves mic access / balance',
    'Video director approves camera sightlines and lighting zones',
    'Orchestra manager approves musician comfort and access',
    'Soloist approves position / sightline to conductor',
    'Venue coordinator approves load-in path and rigging points',
    'All signatures collected on final map',
])

add_small_heading('Setup Procedure — Day 1')
make_table(doc, ['Step', 'Action', 'Owner', 'Est. Time', '\u2713'], [
    ['1', 'Risers placed and leveled', 'Stage crew', '', '\u2610'],
    ['2', 'Chairs placed per map', 'Stage crew', '', '\u2610'],
    ['3', 'Music stands placed + height set', 'Stage crew', '', '\u2610'],
    ['4', 'Conductor podium + stand', 'Stage crew', '', '\u2610'],
    ['5', 'Soloist platform + stand', 'Stage crew', '', '\u2610'],
    ['6', 'Percussion setup / enclosure', 'Perc section', '', '\u2610'],
    ['7', 'Harp positioned', 'Stage crew', '', '\u2610'],
    ['8', 'Orchestra manager walk-through', 'Orch. mgr', '', '\u2610'],
    ['9', 'Conductor walk-through + adjustments', 'Conductor', '', '\u2610'],
    ['10', 'Audio engineer mic placement pass', 'Audio eng.', '', '\u2610'],
    ['11', 'Video director framing pass', 'Video dir.', '', '\u2610'],
    ['12', 'Final walk — all approvers sign off', 'Producer', '', '\u2610'],
], col_widths=[1, 6.5, 3, 2.5, 1.5])

add_small_heading('Riser Diagram')
add_mono_block(doc, '''    [Back Wall]
    ┌────────────┐
    │  Cb / Perc │  Riser 3 (30cm)
    ├────────────┤
    │ Va / Vc    │  Riser 2 (20cm)
    ├────────────┤
    │ Vln I/II   │  Riser 1 (10cm)
    ├────────────┤
    │ Winds/Brass│  Floor
    └────────────┘
    [Soloist / Conductor]''', size=9)

# ═══════════════════════════════════════
# SECTION 5
# ═══════════════════════════════════════
add_page_break()
doc.add_heading('5. Rehearsal Week', level=2)

add_goal_logic(
    'Run a structured rehearsal week that prepares the ensemble for efficient recording sessions.',
    'Every hour of rehearsal saves two hours of recording. But rehearsal is also when the conductor discovers balance issues, the soloist finds tempi, and the engineer dials in the room. The schedule must balance musical preparation with technical discovery.'
)

add_small_heading('Daily Rehearsal Flow')
add_mono_block(doc, '''08:30 — Call (musicians arrive, warm up)
09:00 — Session start (tutti, sectional, or run)
10:30 — Break (union, 15 min)
10:45 — Session resumes
12:30 — Lunch break
13:30 — Afternoon session
15:00 — Break (15 min)
15:15 — Afternoon session resumes
17:00 — Session end (or overtime if approved)''', size=9)

add_small_heading('Rehearsal Notes Log (Daily)')
make_table(doc, ['Day', 'Movement/Passage', 'Issue', 'Resolution', 'Flag for Recording?'], [
    ['Mon', '', '', '', '\u2610Y / \u2610N'],
    ['Tue', '', '', '', '\u2610Y / \u2610N'],
    ['Wed', '', '', '', '\u2610Y / \u2610N'],
    ['Thu', '', '', '', '\u2610Y / \u2610N'],
    ['Fri', '', '', '', '\u2610Y / \u2610N'],
], col_widths=[1.5, 3, 4, 4, 2.5])

add_field('Rundown Template (copy per day)')
add_mono_block(doc, '''Warm-up (10 min) \u2192 Rehearsal block 1 (60 min) \u2192 Break (15)
\u2192 Rehearsal block 2 (60 min) \u2192 Notes (10) \u2192 Run (40)
\u2192 Notes \u2192 Dismiss''', size=9)

add_small_heading('Sectional Schedule')
make_table(doc, ['Day', 'Section', 'Coach/Leader', 'Room', 'Focus'], [
    ['Mon AM', 'Strings', '', '', 'Intonation, bowing consistency'],
    ['Mon PM', 'Winds/Brass', '', '', 'Balance, blend, tuning'],
    ['Tue AM', 'Percussion + Harp', '', '', 'Rhythmic alignment, dynamics'],
    ['Wed PM', 'Tutti (orchestra only)', '', '', 'Full ensemble blend, transitions'],
], col_widths=[2, 2.5, 3, 2.5, 5.5])

add_small_heading('Soloist Rehearsal Integration')
make_table(doc, ['Session', 'Focus', 'Notes'], [
    ['First combined run', 'Tempi agreement, cue points, balance', ''],
    ['Second combined run', 'Expression, phrasing alignment', ''],
    ['Dress rehearsal', 'Full run as live (no stopping)', ''],
    ['Post-dress notes', 'Section pick-ups, adjustments', ''],
], col_widths=[4.5, 6.5, 4])

# ═══════════════════════════════════════
# SECTION 6
# ═══════════════════════════════════════
add_page_break()
doc.add_heading('6. In Between Weekend', level=2)

add_goal_logic(
    'Complete all reconfiguration and prep between rehearsal week and recording week while the orchestra is off.',
    'This is the only gap. The stage must shift from "rehearsal layout" to "recording layout" — chairs may move, audio/video gear gets finalized, and audience ingress/egress gets tested. Every task not done here pushes into recording week, which costs take time.'
)

add_small_heading('Chair Reconfiguration')
make_table(doc, ['Change', 'From', 'To', 'Owner'], [
    ['Chair spacing', 'Rehearsal (loose)', 'Recording (tight grid)', 'Stage crew'],
    ['Stand height', 'Reading height', 'Recording height (fixed)', 'Stage crew'],
    ['Soloist position', 'Rehearsal position', 'Recording + video position', 'Producer'],
    ['Conductor podium', 'Rehearsal position', 'Recording position', 'Producer'],
    ['Aisle width', 'Standard', 'Camera-track ready', 'Stage crew'],
], col_widths=[3.5, 4, 5, 3])

add_small_heading('Door Management (Audience / Guests)')
make_table(doc, ['Item', 'Details', 'Owner'], [
    ['Guest list finalized', '__________________', 'Producer'],
    ['Will-call / check-in process', '__________________', 'PR'],
    ['Seating plan (if applicable)', '__________________', 'Stage mgr'],
    ['House manager assigned', '__________________', 'Producer'],
    ['Signage: entrance, exits, restrooms', '__________________', 'House mgr'],
    ['Program / notes printed', '__________________', 'PR'],
    ['Post-show reception (if any)', '__________________', 'PR'],
], col_widths=[5.5, 7.5, 3])

add_small_heading('Audience Flow Plan')
add_mono_block(doc, '''Arrival \u2192 Will-call check-in \u2192 Seating area (hold until doors open)
\u2192 Doors open \u2192 Audience seated \u2192 Doors close \u2192 Session starts
\u2192 Intermission (if any) \u2192 Session resumes \u2192 End \u2192 Exit''', size=9)

add_small_heading('Equipment Check — Weekend Tasks')
add_checklist(doc, [
    'All microphones tested with cable paths verified',
    'Headphone mix system tested for conductor + soloist',
    'Video cameras rigged, balanced, focus-checked',
    'Video village / control room monitor chain tested',
    'Recording rig: DAW session templates loaded, input mapping verified',
    'Backup recording system running in parallel (mirror)',
    'Timecode sync verified between all A/V systems',
    'Comms system (talkback, cue lights) tested',
    'Power distribution checked — no ground loops, no noisy circuits',
    'Climate control confirmed (orchestra comfort affects takes)',
    'Green room / break area stocked',
    'Parking / load-in access confirmed for recording week',
])

add_field('Weekend Contact (who to call if something breaks)')

# ═══════════════════════════════════════
# SECTION 7
# ═══════════════════════════════════════
add_page_break()
doc.add_heading('7. Recording Week', level=2)

add_goal_logic(
    'Execute the recording sessions efficiently — capture performance takes with minimal retakes, manage session flow, and log everything.',
    'Recording week is the budget driver. Every hour of overtime costs premium rates. The session must run on a tight clock with clear communication: what take we\'re on, what\'s a keeper, what needs a retake, and when to move on.'
)

add_field('Recording Block Template')
add_mono_block(doc, '''09:00 — Call (musicians seated, tuned)
09:15 — Take 1 (movt/section)
09:45 — Playback / discussion
09:55 — Take 2
10:25 — Break
10:40 — Take 3
11:10 — Review / mark keepers
11:30 — Next movement or section
12:30 — Lunch
...''', size=9)

add_small_heading('Daily Session Log')
make_table(doc, ['Take #', 'Movement / Passage', 'Timecode In', 'Timecode Out', 'Keeper?', 'Notes'], [
    ['1', '', '', '', '\u2610', ''],
    ['2', '', '', '', '\u2610', ''],
    ['3', '', '', '', '\u2610', ''],
    ['4', '', '', '', '\u2610', ''],
    ['5', '', '', '', '\u2610', ''],
    ['6', '', '', '', '\u2610', ''],
    ['7', '', '', '', '\u2610', ''],
    ['8', '', '', '', '\u2610', ''],
    ['9', '', '', '', '\u2610', ''],
    ['10', '', '', '', '\u2610', ''],
], col_widths=[1.5, 3.5, 2.5, 2.5, 1.5, 3.5])

add_small_heading('Session Roles — Who Does What During Recording')
make_table(doc, ['Role', 'Responsibility'], [
    ['Recording Producer', 'Runs the room, calls takes, speaks to conductor/soloist, marks keepers'],
    ['Recording Engineer', 'Operates DAW, manages levels, monitors input, tracks timecode'],
    ['Assistant Engineer', 'Cable management, mic adjustments, talkback, log sheets'],
    ['Video Director', 'Camera switching/recording, sync verification, shot logging'],
    ['Conductor', 'Musical direction, tempo, cueing orchestra, take feedback'],
    ['Soloist', 'Performance, take feedback, retake requests'],
    ['Orchestra Manager', 'Musician welfare, break enforcement, overtime tracking'],
    ['Librarian', 'Score copies, page turns, part corrections on the fly'],
    ['Stage Manager', 'Stage changes, musician entry/exit, green room'],
], col_widths=[3.5, 12])

add_small_heading('Take Marking Convention')
add_mono_block(doc, '''KEEPER  — green highlighter on log sheet, green tape on DAW session
MAYBE   — yellow, needs comparison in patching
REJECT  — red, technical or performance issue
PATCH   — blue, good performance but needs patch from another take
PROBLEM — circled red, note why''', size=9)

add_small_heading('End-of-Day Routine')
add_checklist(doc, [
    'Daily session log reviewed and filed',
    'Backup drives swapped / verified',
    'Next day\'s session plan confirmed with conductor + soloist',
    'Equipment powered down properly (or left on if overnight)',
    'Cameras/media backed up',
    'Clock signed (union)',
    'Tomorrow\'s call time confirmed and posted',
])

# ═══════════════════════════════════════
# SECTION 8
# ═══════════════════════════════════════
add_page_break()
doc.add_heading('8. Video Adjustment', level=2)

add_goal_logic(
    'Position people, chairs, and microphones for optimal video composition without compromising audio quality.',
    'Video and audio have competing needs. Audio wants mics close and visible; video wants them gone. Chairs that work for audio balance may block camera sightlines. This section captures the compromises and ensures every position is intentional and approved.'
)

add_small_heading('People Positioning — Camera Map')
add_mono_block(doc, '''                          [CONDUCTOR]
                              |
                     [SOLOIST — centered]
                              |
   ┌──────────────┐──────────┼──────────┐──────────────┐
   │   Vln I      │   Vln II  │  Vla     │   Vc         │
   │ Cam A (wide) │   Cam B   │  Cam C   │  Cam D (flank)│
   └──────────────┴──────────┴──────────┴──────────────┘

Key:
WIDE   — full stage + conductor
SOLO   — soloist close-up, 45° from conductor side
COND   — conductor face, behind-orchestra angle''', size=8)

make_table(doc, ['Camera', 'Shot', 'Lens', 'Position', 'Notes'], [
    ['Cam A (Wide)', 'Full stage', '__________________', '__________________', ''],
    ['Cam B (Solo)', 'Soloist close', '__________________', '__________________', ''],
    ['Cam C (Cond)', 'Conductor', '__________________', '__________________', ''],
    ['Cam D (Flank)', 'Strings/LH side', '__________________', '__________________', ''],
    ['Cam E (Optional)', 'Winds/Brass', '__________________', '__________________', ''],
], col_widths=[3, 2.5, 3, 3.5, 3])

add_small_heading('Chair Adjustment for Camera Sightlines')
make_table(doc, ['Section', 'Sightline Issue', 'Adjustment', 'Approved By'], [
    ['Vln I (outside)', 'Blocks soloist', 'Shift 15cm right', ''],
    ['Vln II (inside)', 'Blocks conductor angle', 'Lower riser or shift', ''],
    ['Harp', 'Reflection in solo cam', 'Rotate 10°', ''],
    ['Percussion', 'Camera C blocked', 'Riser height adjust', ''],
    ['Timpani', 'Rear cam sightline', 'Shift 20cm left', ''],
], col_widths=[3, 4, 3.5, 3.5])

add_small_heading('Microphone Placement — Video-Conscious')
make_table(doc, ['Mic', 'Purpose', 'Audio Ideal', 'Video-Adjusted', 'Approved'], [
    ['Main pair (ORFT/AB)', 'Orchestra stereo', '2m above conductor', 'Raised to 2.5m', '\u2610'],
    ['Spot — Soloist', 'Solo close', '60cm, 45°', '80cm, 30° (out of frame)', '\u2610'],
    ['Spot — Principal winds', 'Woodwind clarity', 'Over bell', 'Side angle, out of frame', '\u2610'],
    ['Spot — Harp', 'Harp clarity', 'Above soundboard', 'Below frame line', '\u2610'],
    ['Ambient / hall', 'Room sound', '10m back, 4m high', 'No change', '\u2610'],
    ['Conductor cam mic', 'Conductor audio (sync)', 'Lavalier', 'Hidden under collar', '\u2610'],
], col_widths=[3, 2.5, 2.5, 3.5, 2.5])

add_small_heading('Lighting Zones')
make_table(doc, ['Zone', 'Purpose', 'Fixture', 'Color Temp', 'Notes'], [
    ['Soloist', 'Primary focus', 'Key + fill', '3200K', 'No shadow on face'],
    ['Conductor', 'Readable face', 'Front wash', '3200K', 'Must see baton'],
    ['Orchestra', 'Uniform coverage', 'General wash', '3200K', 'Avoid hot spots'],
    ['Background', 'Depth', 'Back light', '4000K', 'Separate from backdrop'],
], col_widths=[2.5, 3, 3, 2.5, 4.5])

add_small_heading('Video Adjustment — Final Approval')
add_checklist(doc, [
    'All camera positions locked and marked on floor',
    'Chair positions adjusted per sightline map',
    'Mic positions adjusted (audio balance re-checked after each move)',
    'Lighting levels set for video (not too hot, not too cool)',
    'White balance set on all cameras',
    'Timecode jammed across all cameras and audio recorder',
    'Recording engineer confirms audio acceptable at all mic positions',
    'Video director confirms all shots framed',
    'Conductor confirms sightlines to orchestra are not obstructed',
    'Producer signs off on final positions',
])

# ═══════════════════════════════════════
# SECTION 9
# ═══════════════════════════════════════
add_page_break()
doc.add_heading('9. Recording — Session Procedure', level=2)

add_goal_logic(
    'Define the exact procedure for every recording session: who speaks, when, and how takes are managed.',
    'Recording is high-pressure and expensive. Ambiguity about when to roll, who calls "cut," or whether a take is a keeper burns time and creates tension. A documented procedure means the room runs itself.'
)

add_small_heading('Recording Procedure — Step by Step')
add_mono_block(doc, '''PRE-ROLL (Producer or Engineer):
1. "Standby for take [number], [movement/section]."
2. "Recording rolling."  (confirm record light / DAW transport)
3. "Slate."              (engineer slates take number)

ROLLING:
4. Conductor gives downbeat.
5. Performance proceeds. No one speaks in the room except in emergency.
6. Engineer watches levels, video rolls.

POST-ROLL:
7. Conductor cuts off.
8. Engineer stops recording.  Note timecode.
9. Producer (only) speaks first: "How did that feel?"
   — Conductor responds.
   — Soloist responds (if applicable).
   — Engineer gives technical notes (if any: distortion, noise, dropout).
10. Producer decides:
    — "Moving on."        (keeper, next section)
    — "One more."          (good but needs retry)
    — "Patch that in [take X]." (this take has a usable passage)
    — "Let's hear it back."     (playback requested)''', size=8.5)

add_small_heading('Comms Protocol')
make_table(doc, ['Channel / Method', 'Used For', 'Who Speaks'], [
    ['Talkback (control room \u2192 hall)', 'Producer to room, take calls', 'Producer only'],
    ['Headphones (conductor)', 'Producer direct to conductor', 'Producer \u2194 Conductor'],
    ['Headphones (soloist)', 'Producer direct to soloist', 'Producer \u2194 Soloist'],
    ['Cue lights (hall)', 'Visual signal: recording active', 'Engineer controls'],
    ['Private channel (engineer \u2194 producer)', 'Technical discussion mid-take', 'Eng. \u2194 Producer'],
    ['Private channel (video)', 'Camera switching, shot calls', 'Video director'],
], col_widths=[4.5, 5, 5])

add_small_heading('Slate Convention')
add_mono_block(doc, '''FORMAT:  [Date]_[Movement]_[TakeNumber]
EXAMPLE: 0721_Mvmt1_T03

Audio files named this way in DAW.
Video files named the same + _V at the end.
Log sheet matches exactly. No variations.''', size=9)

add_small_heading('Break Enforcement')
add_checklist(doc, [
    'Union break required every [X] minutes — track from session start',
    'Producer watches the clock; orchestra manager enforces the break',
    'No recording during break — full silence (engineer may keep rig live)',
    'Break start / end announced via talkback',
    'If mid-take at break time: finish the take, then break',
])

add_small_heading('Emergency Recording Procedure')
add_mono_block(doc, '''TECHNICAL FAILURE (power, hard drive, mic failure):
1. Engineer calls "Stop."  Immediate. No hesitation.
2. Note timecode of failure.
3. Producer announces: "Technical hold. Stay seated."
4. Engineer diagnoses and resolves.
5. Restart: "Standby, rolling from [timecode of last good frame]."
6. If quick fix: retake from start of movement.
   If extended fix (>15 min): announce break.

MEDICAL EMERGENCY:
1. Engineer stops recording.
2. Stage manager handles medical.
3. Producer announces: "Hold. Medical situation."
4. Musicians stay seated unless directed otherwise.
5. Resume only after all-clear from stage manager.''', size=8.5)

# ═══════════════════════════════════════
# SECTION 10
# ═══════════════════════════════════════
add_page_break()
doc.add_heading('10. Patching (Post-Production)', level=2)

add_goal_logic(
    'Compile the best takes into a seamless master — comping audio takes, syncing video, and delivering the final product.',
    'Great recording sessions produce 70% keepers and 30% patches. Patching is where the final performance is assembled. Without a clear system, the editor guesses at which take the conductor preferred, and the producer spends days reconstructing intent.'
)

add_small_heading('Patching Workflow')
add_mono_block(doc, '''PHASE 1 — AUDIO COMP (Editor)
Session review \u2192 Import marked keepers \u2192 Comp per movement
\u2192 Crossfade edits \u2192 Check for edit clicks \u2192 Rough mix

PHASE 2 — VIDEO SYNC (Video Editor)
Import rough audio \u2192 Sync video takes to comp \u2192 Cut video per audio
\u2192 Multi-cam switch \u2192 Color grade \u2192 Final sync check

PHASE 3 — REVIEW (Producer + Conductor + Soloist)
Listen / watch pass \u2192 Mark revision requests \u2192 Return to editor
\u2192 Revise \u2192 Final approval

PHASE 4 — DELIVERY
Master audio export (48kHz/24bit minimum) \u2192 Master video export
\u2192 Deliverables per distribution agreement''', size=8.5)

add_small_heading('Keeper / Patch Log')
make_table(doc, ['Movement', 'Primary Take', 'Patch Source', 'Edit Point', 'Status'], [
    ['Mvt 1 — Intro', '', '', '', '\u2610'],
    ['Mvt 1 — Exposition', '', '', '', '\u2610'],
    ['Mvt 1 — Development', '', '', '', '\u2610'],
    ['Mvt 1 — Recapitulation', '', '', '', '\u2610'],
    ['Mvt 1 — Coda', '', '', '', '\u2610'],
    ['Mvt 2 — Opening', '', '', '', '\u2610'],
    ['Mvt 2 — Theme', '', '', '', '\u2610'],
    ['Mvt 2 — Middle section', '', '', '', '\u2610'],
    ['Mvt 2 — Close', '', '', '', '\u2610'],
    ['Mvt 3 — Opening', '', '', '', '\u2610'],
    ['Mvt 3 — Development', '', '', '', '\u2610'],
    ['Mvt 3 — Cadenza', '', '', '', '\u2610'],
    ['Mvt 3 — Finale', '', '', '', '\u2610'],
], col_widths=[3.5, 2.5, 2.5, 3, 1.5])

add_small_heading('Deliverables Checklist')
make_table(doc, ['Deliverable', 'Format', 'Spec', 'Delivered?', 'Recipient'], [
    ['Stereo audio master', '.wav', '48kHz/24bit', '\u2610', ''],
    ['Multitrack audio stems', '.wav', 'Per mic group', '\u2610', ''],
    ['Video master (full)', 'ProRes / H.264', '4K / 1080p', '\u2610', ''],
    ['Video edit (highlights)', 'H.264', '1080p', '\u2610', ''],
    ['Session archive', 'DAW project', 'All takes + comps', '\u2610', ''],
    ['Metadata / cue sheet', 'PDF', 'Track timings, composers', '\u2610', ''],
    ['Cover art / credits', 'PDF / TIFF', '', '\u2610', ''],
], col_widths=[4.5, 2.5, 3.5, 2, 3])

add_small_heading('Patching Approval Chain')
add_checklist(doc, [
    'Rough comp delivered to producer (audio only)',
    'Producer reviews and marks changes',
    'Conductor listens to rough comp and approves / requests changes',
    'Soloist listens to rough comp and approves / requests changes',
    'Final comp locked and signed by producer',
    'Video edit delivered to producer',
    'Video reviewed and approved',
    'Master exports QC\'d (full listen/watch-through)',
    'Deliverables sent to recipient',
])

add_small_heading('Revision Tracking')
make_table(doc, ['Revision #', 'Date', 'Requested By', 'Change', 'Status'], [
    ['', '', '', '', ''],
    ['', '', '', '', ''],
    ['', '', '', '', ''],
    ['', '', '', '', ''],
], col_widths=[2, 2, 3, 4.5, 3.5])

# ═══════════════════════════════════════
# SECTION 11
# ═══════════════════════════════════════
add_page_break()
doc.add_heading('11. Unexpected Changes — Contingency Plans', level=2)

add_goal_logic(
    'Anticipate the most likely failures and have a pre-agreed response for each.',
    'The production that scrambles in crisis is the one that wastes money and produces a rushed product. A written contingency for each common failure means the team executes the plan instead of inventing one under pressure.'
)

add_small_heading('Common Failure Modes — Pre-Planned Responses')
make_table(doc, ['Failure', 'Probability', 'Impact', 'Response', 'Owner'], [
    ['Soloist cancels (illness)', 'Low', 'Critical', 'Identify backup dates before project start. If day-of: record orchestral-only takes, schedule soloist pick-up session.', 'Producer'],
    ['Conductor unavailable', 'Low', 'Critical', 'Designated assistant conductor prepped. Or: postpone to backup date.', 'Producer'],
    ['Key musician cancels (principal)', 'Medium', 'High', 'Sub list confirmed before project. Sub prepped with parts in advance.', 'Orch. mgr'],
    ['Venue double-booked / unavailable', 'Low', 'Critical', 'Venue contract includes penalty clause for double-booking. Backup venue scouted before contract signed.', 'Producer'],
    ['Recording rig failure (primary)', 'Medium', 'High', 'Mirror recording system running at all times. If primary fails, switch to mirror. No lost takes.', 'Audio eng'],
    ['Hard drive failure', 'Low', 'High', 'Triple redundancy: internal + external + cloud sync. Daily verification.', 'Audio eng'],
    ['Hard drive failure', 'Low', 'High', 'Triple redundancy: internal + external + cloud sync. Daily verification.', 'Audio eng'],
    ['Power outage', 'Low', 'Critical', 'UPS on all recording gear. Venue generator backup confirmed. Session pauses until power stable.', 'Venue coord'],
    ['Noise contamination (siren, HVAC)', 'Medium', 'Medium', 'Pause take immediately. Note timecode. Retake from nearest clean entry point. HVAC schedule confirmed with venue.', 'Engineer'],
    ['Lighting failure (video)', 'Low', 'Medium', 'Spare lighting kit on-site. Video switches to available light + grade in post.', 'Video dir'],
    ['Audience disruption', 'Medium', 'Low', 'House manager handles. Recording continues unless prolonged. If prolonged: pause, address, resume.', 'House mgr'],
    ['Union overtime limit reached', 'Medium', 'Medium', 'Pre-agreed overtime budget. If exceeded, stop recording. Schedule overflow session.', 'Producer'],
    ['Music part error (wrong note, missing page)', 'Medium', 'Medium', 'Librarian on-site with master score. Correction made between takes. If major: schedule pick-up.', 'Librarian'],
], col_widths=[4.5, 1.5, 2, 6.5, 2.5])

add_small_heading('Backup Date Planning')
add_field('Recording Window (primary)')
add_field('Backup Window 1')
add_field('Backup Window 2')
add_field('Soloist Backup Availability')
add_field('Conductor Backup Availability')

add_small_heading('Decision Tree — Cancel vs. Continue')
add_mono_block(doc, '''IS THE ESSENTIAL PERSON / SYSTEM AVAILABLE?
    YES \u2192 Continue
    NO  \u2192 CAN WE WORK AROUND?
             YES \u2192 Workaround (e.g., record other movements first)
             NO  \u2192 CAN WE RESCHEDULE WITHIN BUDGET?
                      YES \u2192 Rebook
                      NO  \u2192 CANCEL''', size=9)

add_small_heading('Emergency Contact Tree')
add_mono_block(doc, '''PRODUCER \u2192 makes all go/no-go decisions
  \u2192 if unreachable: Artistic Director
  \u2192 if unreachable: Orchestra Manager

TECHNICAL EMERGENCY (audio/video/lighting):
  Engineer \u2192 Producer \u2192 Go/No-Go

MEDICAL EMERGENCY:
  Stage Manager \u2192 Venue First Aid \u2192 Producer (informational)
  Conductor and Producer decide on recording pause/resume

MUSICIAN ABSENCE:
  Orchestra Manager \u2192 Producer \u2192 Substitute confirmed''', size=9)

add_field('Producer — Go/No-Go Decision Authority')
add_field('Budget Contingency (dollars or percentage)')

# ═══════════════════════════════════════
# POST-EVENT
# ═══════════════════════════════════════
add_page_break()
doc.add_heading('Post-Event Notes', level=2)

p = doc.add_paragraph()
r = p.add_run('Complete after delivery and file this book with notes clipped inside the front cover. The book improves every project.')
r.font.italic = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

add_small_heading('What ran late?')
p = doc.add_paragraph()
r = p.add_run('_' * 70)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

add_small_heading('What broke / failed?')
p = doc.add_paragraph()
r = p.add_run('_' * 70)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

add_small_heading('What confused the crew?')
p = doc.add_paragraph()
r = p.add_run('_' * 70)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

add_small_heading('What to change next time?')
p = doc.add_paragraph()
r = p.add_run('_' * 70)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

add_small_heading('Budget actual vs. estimate')
p = doc.add_paragraph()
r = p.add_run('_' * 70)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

add_small_heading('New sections to add / remove?')
p = doc.add_paragraph()
r = p.add_run('_' * 70)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

output_path = '/Users/patricktrang/opencode-2607/concerto-promptbook.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
