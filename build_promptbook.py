from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Style Helpers ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for level in [1, 2, 3]:
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    if level == 1:
        h.font.size = Pt(22)
        h.font.bold = True
    elif level == 2:
        h.font.size = Pt(15)
        h.font.bold = True
    elif level == 3:
        h.font.size = Pt(11)
        h.font.bold = True
        h.font.color.rgb = RGBColor(0x40, 0x40, 0x60)

def add_page_break():
    doc.add_page_break()

def add_goal_logic(goal, logic):
    p = doc.add_paragraph()
    r = p.add_run('Goal: ')
    r.bold = True
    r.font.size = Pt(10.5)
    r = p.add_run(goal)
    r.font.size = Pt(10.5)
    p2 = doc.add_paragraph()
    r = p2.add_run('Logic: ')
    r.bold = True
    r.font.size = Pt(10.5)
    r = p2.add_run(logic)
    r.font.size = Pt(10.5)

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1A1A2E')
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_shading(cell, 'F0F0F5')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def add_checklist(doc, items):
    for item in items:
        p = doc.add_paragraph()
        r = p.add_run(f'\u2610  {item}')
        r.font.size = Pt(10)

def add_field(label, width_cm=8):
    p = doc.add_paragraph()
    r = p.add_run(f'{label}: ')
    r.bold = True
    r.font.size = Pt(10)
    r = p.add_run('_' * int(width_cm * 3))
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

def add_mono_block(doc, text, size=8.5):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(size)

def add_small_heading(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

# ═══════════════════════════════════════
# SECTION 1 — Cover + Event Overview (combined)
# ═══════════════════════════════════════

# Title block (compact)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('DANCE COMPETITION EVENT PROMPT BOOK')
r.font.size = Pt(18)
r.bold = True
r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(6)
r = p.add_run('─' * 50)
r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

# Meta info inline on one line style
meta_items = [
    ('Event', '[Event Name]'),
    ('Date(s)', '[Date(s)]'),
    ('Venue', '[Venue]'),
    ('Version', '[Draft / Final]'),
]
meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_p.paragraph_format.space_after = Pt(8)
for i, (label, val) in enumerate(meta_items):
    r = meta_p.add_run(f'{label}: ')
    r.bold = True
    r.font.size = Pt(9.5)
    r = meta_p.add_run(val)
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    if i < len(meta_items) - 1:
        r = meta_p.add_run('    ')
        r.font.size = Pt(9.5)

doc.add_heading('1. Event Overview — How to Use This Book', level=2)

add_goal_logic(
    'Explain the sections of the prompt book so anyone on the team can navigate it.',
    'The book must teach itself. A new volunteer or returning team member should open this page, understand what each section contains and when to use it, without needing you there to explain. If the book only makes sense to the person who wrote it, it fails the moment that person is busy putting out fires.'
)

make_table(doc, ['#', 'Section', 'When You Need It'], [
    ['1', 'Event Overview', 'Cover page + how to navigate the book'],
    ['2', 'Task Split + Schedule Requirements', 'Planning — who does what, building the timeline'],
    ['3', 'Policies to Print & Post', 'During venue setup'],
    ['4', 'Registration + Music Collection + Deadlines', 'At competitor intake through submission cutoff'],
    ['5', 'Order Check + Prop Spiking', 'On verification day at the theater'],
    ['6', 'Emcees Script', 'Script prep / during show'],
    ['7', 'Judge Tables', 'Venue setup before competition days'],
    ['8', 'Preliminaries Preparation', 'Day before or morning of prelims'],
    ['9', 'Preliminaries', 'Running the day'],
    ['10', 'Judge Review Sheets & Decision', 'During tabulation'],
    ['11', 'Semi Finals Preparation', 'After prelims results'],
    ['12', 'Semi Finals', 'Running the day'],
    ['13', 'Finals Prep + Awards Prep', 'After semis results, before finals day'],
    ['14', 'Finals', 'Running finals day'],
    ['15', 'Awards', 'Running the ceremony'],
], col_widths=[1, 6, 8.5])

# ═══════════════════════════════════════
# SECTION 2
# ═══════════════════════════════════════
add_page_break()
doc.add_heading('2. Task Split + Schedule Requirements', level=2)

add_goal_logic(
    'List every needed task with one owner per task, then document the hard constraints and iterations that build the schedule.',
    'Live events fail on ownership gaps \u2014 "I thought you had it." One name per task kills ambiguity. The first schedule draft is always wrong; recording constraints and each attempt means the final version is explainable.'
)

make_table(doc, ['Task', 'Owner', 'Backup', 'Unique Challenge / Notes'], [
    ['Registration & Intake', '', '', ''],
    ['Music Collection', '', '', ''],
    ['Order Check Coordination', '', '', ''],
    ['Prop Spiking Lead', '', '', ''],
    ['Sound / Audio Board', '', '', ''],
    ['Lighting Board', '', '', ''],
    ['Stage Management', '', '', ''],
    ['Emcee', '', '', ''],
    ['House Management', '', '', ''],
    ['Judging Coordinator', '', '', ''],
    ['Tabulation / Scoring', '', '', ''],
    ['Awards Setup', '', '', ''],
    ['Photography / Video', '', '', ''],
    ['Volunteer Coordination', '', '', ''],
    ['Venue Liaison', '', '', ''],
    ['Emergency / Medical', '', '', ''],
], col_widths=[4, 3, 3, 5.5])

add_small_heading('Core Teams')
p = doc.add_paragraph()
r = p.add_run('Front-of-House: ')
r.bold = True
r.font.size = Pt(9.5)
r = p.add_run('Registration, house management, audience flow    ')
r.font.size = Pt(9.5)
r = p.add_run('Backstage: ')
r.bold = True
r.font.size = Pt(9.5)
r = p.add_run('Stage management, props, runners, green room    ')
r.font.size = Pt(9.5)
r = p.add_run('Production: ')
r.bold = True
r.font.size = Pt(9.5)
r = p.add_run('Sound, lights, cues, emcee\n')
r.font.size = Pt(9.5)
r = p.add_run('Judging & Scoring: ')
r.bold = True
r.font.size = Pt(9.5)
r = p.add_run('Judge wrangling, sheets, tabulation    ')
r.font.size = Pt(9.5)
r = p.add_run('Logistics: ')
r.bold = True
r.font.size = Pt(9.5)
r = p.add_run('Setup, teardown, awards, venue communication')
r.font.size = Pt(9.5)

add_small_heading('Hard Constraints')
make_table(doc, ['Constraint', 'Value'], [
    ['Venue load-in time', '__________________'],
    ['Venue curfew (must be out by)', '__________________'],
    ['Number of acts registered', '__________________'],
    ['Average act length (minutes)', '__________________'],
    ['Acts per heat', '__________________'],
    ['Transition time between heats', '__________________'],
    ['Judge deliberation per round', '__________________'],
    ['Lunch break duration', '__________________'],
    ['Awards ceremony duration', '__________________'],
], col_widths=[6, 9.5])

add_small_heading('Schedule Attempts Log')
make_table(doc, ['Attempt #', 'Date', 'Key Change', 'Why It Changed'], [
    ['1', '', '(Initial draft)', ''],
    ['2', '', '', ''],
    ['3', '', '', ''],
    ['Final', '', '', ''],
], col_widths=[2, 2.5, 5.5, 5.5])

add_small_heading('Organization Steps')
add_checklist(doc, [
    'Gather registration data (act count, categories, styles)',
    'Estimate average act length + heat duration',
    'Calculate total runtime (acts \u00d7 length + transitions + breaks)',
    'Cross-check against venue available hours',
    'Identify bottlenecks (crowded categories, slow transitions)',
    'Build buffer (15% minimum per session)',
    'Share first draft with production team \u2192 revise \u2192 log \u2192 lock',
])

# ═══════════════════════════════════════
# SECTION 3
# ═══════════════════════════════════════
add_page_break()
doc.add_heading('3. Policies to Print and Post', level=2)

add_goal_logic(
    'Have every enforceable rule on paper and physically posted at the venue before the event starts.',
    'A policy settles disputes only if it exists before the dispute. Post each where its problem happens \u2014 backstage rules at the backstage door, music rules at check-in. A posted policy is enforcement; a verbal one is a debate.'
)

add_checklist(doc, [
    'Backstage Rules \u2014 No parents backstage. Dancers only. Quiet zones. No phones during performance.',
    'Music Rules \u2014 Accepted formats. No profanity. No unlicensed tracks. Late submission penalty.',
    'Competitor Conduct \u2014 Dress code. Check-in deadlines. Judging protest procedure.',
    'Audience Rules \u2014 No flash photography. No moving during acts. No saving seats.',
    'Lost Child / Lost Item Protocol \u2014 Where to go. Who to ask.',
    'Emergency Procedures \u2014 Fire exits. Assembly point. Medical station.',
])

add_small_heading('Posting Location Map')
make_table(doc, ['Policy', 'Location(s)'], [
    ['Backstage Rules', 'Backstage entrance, green room, dressing rooms'],
    ['Music Rules', 'Check-in table, audio booth'],
    ['Competitor Conduct', 'Registration desk, backstage bulletin board'],
    ['Audience Rules', 'Box office, house doors, program insert'],
    ['Lost Child / Lost Item', 'Box office, front-of-house table'],
    ['Emergency Procedures', 'Every exit, green room, lobby'],
], col_widths=[4.5, 11])

# ═══════════════════════════════════════
# SECTION 4
# ═══════════════════════════════════════
add_page_break()
doc.add_heading('4. Registration + Music Collection + Deadlines', level=2)

add_goal_logic(
    'One intake flow that catches errors early: enforce policies, collect titles, draw order numbers, collect music from captains, and set a deadline that compresses last-minute changes.',
    'An error caught at intake costs two minutes \u2014 caught at order check it stalls a line. Chasing dozens of individual competitors is chaos; chasing captains is manageable. The deadline forces bulk submission early so order check becomes verification, not data entry.'
)

add_small_heading('Registration Check-in Flow')
add_mono_block(doc, 'Competitor arrives  \u2192  Policy acknowledgment signed  \u2192  Naming format verified\n\u2192  Order number drawn  \u2192  Title collected  \u2192  Dance captain form issued', size=9)

add_field('Naming Format Rule')
p = doc.add_paragraph()
r = p.add_run('(Example: STUDIO_CATEGORY_ACTNUMBER_DANCERNAME)')
r.font.size = Pt(8.5)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
r.font.italic = True

add_field('Order Number Drawing Method')
p = doc.add_paragraph()
r = p.add_run('(e.g., random draw per category, first-come-first-served)')
r.font.size = Pt(8.5)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
r.font.italic = True

add_small_heading('Music Collection Station')
make_table(doc, ['Item', 'Details'], [
    ['Location', '__________________'],
    ['Equipment', 'Laptop / USB drive / File checker'],
    ['File format accepted', '__________________'],
    ['Naming convention enforced', '__________________'],
], col_widths=[5, 10.5])

add_small_heading('Captain Check-in Flow')
add_mono_block(doc, 'Captain arrives  \u2192  Presents files  \u2192  Staff copies + reads back confirmation\n\u2192  Captain signs off  \u2192  Deadline reminder given', size=9)

add_small_heading('Music Collection Log')
make_table(doc, ['Captain', 'Class', '# Files', 'Names OK?', 'Format OK?', 'Staff', 'Date'], [
    ['', '', '', '\u2610Y/\u2610N', '\u2610Y/\u2610N', '', ''],
    ['', '', '', '\u2610Y/\u2610N', '\u2610Y/\u2610N', '', ''],
    ['', '', '', '\u2610Y/\u2610N', '\u2610Y/\u2610N', '', ''],
    ['', '', '', '\u2610Y/\u2610N', '\u2610Y/\u2610N', '', ''],
], col_widths=[2.5, 2, 1.5, 2, 2, 2.5, 2.5])

add_small_heading('Deadline & Change Tracking')
add_field('Official Deadline (date / time)')

make_table(doc, ['Date', 'Time', 'Competitor', 'What Changed', 'Approved By'], [
    ['', '', '', '', ''],
    ['', '', '', '', ''],
    ['', '', '', '', ''],
], col_widths=[2.5, 2, 3.5, 4.5, 3])

add_small_heading('Deadline Day Staffing')
make_table(doc, ['Role', 'Details'], [
    ['Music intake staff (people)', '__________________'],
    ['Hours covered', '__________________'],
    ['After-hours contact', '__________________'],
], col_widths=[5.5, 10])

# ═══════════════════════════════════════
# SECTION 5
# ═══════════════════════════════════════
add_page_break()
doc.add_heading('5. Order Check + Prop Spiking', level=2)

add_goal_logic(
    'Run two parallel verification streams on the same day \u2014 competitors confirm order + music in the audio booth while a separate crew spikes props on stage for SM notes.',
    'Every competitor is already in the building for order check. This is free stage access, your scarcest resource. Running both simultaneously means the SM has documented prop positions and every track has been heard through the actual system before show day.'
)

add_small_heading('Order Check Station Setup')
make_table(doc, ['Station', 'Location'], [
    ['Audio check station', '__________________'],
    ['Order confirmation table', '__________________'],
    ['Runner path between stations', '__________________'],
], col_widths=[5.5, 10])

add_small_heading('Competitor Flow')
add_mono_block(doc, 'Arrive  \u2192  Confirm name, act #, category  \u2192  Walk to audio booth\n\u2192  Hear track played  \u2192  Approve or flag issue  \u2192  Sign off', size=9)

add_small_heading('Order Check Confirmation Form')
add_mono_block(doc, '''Competitor: ___________  Act#: ________  Category: ________
\u2610 Order confirmed          \u2610 Track approved
\u2610 Issue: ______________________________
   Resolution: ______________________________
Signature: _____________  Audio: _____________  Date: ______________''', size=9)

add_small_heading('Issue Tracking')
make_table(doc, ['Competitor', 'Issue', 'Resolution', 'Resolved By'], [
    ['', '', '', ''],
    ['', '', '', ''],
], col_widths=[3, 4.5, 4.5, 3.5])

add_small_heading('Prop Spiking')
make_table(doc, ['Role', 'Name'], [
    ['Lead', '__________________'],
    ['Stage Manager', '__________________'],
    ['Assistants', '__________________'],
], col_widths=[4.5, 11])

add_small_heading('Prop Spiking Schedule (parallel to order check)')
make_table(doc, ['Time', 'Competitor', 'Prop', 'Stage Position', 'SM Init.'], [
    ['', '', '', '', ''],
    ['', '', '', '', ''],
    ['', '', '', '', ''],
], col_widths=[2.5, 3, 3.5, 4, 2])

add_small_heading('Stage Map')
add_mono_block(doc, '''        [Upstage]
          UL    UC    UR
          ML    MC    MR
          DL    DC    DR
      [Downstage / Audience]''', size=9.5)

p = doc.add_paragraph()
r = p.add_run('Key:  \u2b1c = Set piece    \u25b3 = Hand prop drop    \u25a1 = Floor marking    ~ = Fabric')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# ═══════════════════════════════════════
# SECTION 6
# ═══════════════════════════════════════
add_page_break()
doc.add_heading('6. Emcees Script', level=2)

add_goal_logic(
    'Write the majority of the script section by section and establish the general rules governing it.',
    'Rules first \u2014 name pronunciation, sponsor phrasing, tone, dead-time fill \u2014 then the section scripts. A scripted skeleton means day-of changes are small edits, not improvisation, and the emcee stays consistent across all sessions.'
)

add_small_heading('General Script Rules')
for i in range(5):
    p = doc.add_paragraph()
    r = p.add_run(f'{i+1}.  ___________________________________________________________')
    r.font.size = Pt(10)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
r = p.add_run('(Examples: "Every sponsor mention uses full name, never abbreviated." / "Competitor names read as [First Last], no nicknames unless approved." / "Dead time = read the next category description.")')
r.font.size = Pt(8.5)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
r.font.italic = True

add_small_heading('Script Sections to Write')
add_checklist(doc, [
    'Welcome / Housekeeping (doors, phones, emergency exits)',
    'Sponsor acknowledgments',
    'Judge introductions',
    'Category lead-ins (read before each section)',
    'Intermission announcements',
    'Results / advancement announcements',
    'Awards ceremony script',
    'Closing remarks',
])

add_small_heading('Script Template Format')
add_mono_block(doc, '''\u255d\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550
SECTION: [Category Name / Segment]
CUES: [Light cue, sound cue, SM cue]
\u255d\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550
EMCEE:
[Scripted text here. Placeholders in [BRACKETS].]

[Optional: ad-lib limits / notes]

\u2192 Cue: [Next action]''', size=9)

# ═══════════════════════════════════════
# SECTION 7
# ═══════════════════════════════════════
add_page_break()
doc.add_heading('7. Judge Tables', level=2)

add_goal_logic(
    'Define the complete judge table setup: position, tablecloths, lighting, notation folders, pencils, snacks.',
    'A judge looking for a pencil or fighting glare is a judge not watching the dancer. Standardizing the full kit makes setup a checklist, and locking table positions protects sightlines from day-of rearranging.'
)

add_small_heading('Position Diagram')
add_mono_block(doc, '''        [STAGE]
             \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
             \u2502             \u2502
             \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518
          [Audience / Judges]
           J1    J2    J3    J4''', size=9.5)

make_table(doc, ['Measurement', 'Value'], [
    ['Distance from stage edge', '__________________'],
    ['Angle from center', '__________________'],
    ['Table dimensions', '__________________'],
    ['Number of judges', '__________________'],
], col_widths=[5.5, 10])

add_small_heading('Per-Table Kit Checklist')
make_table(doc, ['Item', 'Qty', '\u2713'], [
    ['Table (standard height)', '1', '\u2610'],
    ['Tablecloth (color: _______)', '1', '\u2610'],
    ['Desk lamp (warm light, no stage glare)', '1', '\u2610'],
    ['Power strip / extension cord', '1', '\u2610'],
    ['Judge folder (scoring sheets, competitor list, rubric, pen)', '1', '\u2610'],
    ['Pencils (sharpened, no mechanical)', '3', '\u2610'],
    ['Pencil sharpener', '1', '\u2610'],
    ['Bottled water', '1', '\u2610'],
    ['Snack (quiet, no crinkle wrappers)', '1', '\u2610'],
    ['Sign: "JUDGE \u2014 DO NOT DISTURB DURING ROUTINES"', '1', '\u2610'],
], col_widths=[9, 1.5, 1.5])

add_small_heading('Setup Checklist')
add_checklist(doc, [
    'Tables positioned, leveled, sightlines verified \u2014 no obstruction between judge and full stage',
    'Tablecloths clean, draping evenly',
    'Lamps plugged in, angled to light the paper (not the stage)',
    'Each folder stocked with correct scoring sheets for this session',
    'Extra pencils + sharpener on a side table (not the judge table)',
    'Water and snack placed',
    'Signs posted',
])

# ═══════════════════════════════════════
# SECTION 8
# ═══════════════════════════════════════
add_page_break()
doc.add_heading('8. Preliminaries Preparation', level=2)

add_goal_logic(
    'Cover what must be posted, what must be announced, and the final check before prelim day starts.',
    'This page catches the unowned tasks \u2014 signage, drafted announcements, printed files, current master file. The "last final check" is a physical walk of the venue with this book in hand.'
)

add_field('Prelims Date')

add_small_heading('What to Post')
add_checklist(doc, [
    'Competition order / act list (backstage + lobby)',
    'Category schedule with estimated times (lobby)',
    'Backstage rules (green room, dressing rooms)',
    'Emergency exits and assembly point (all areas)',
    'Lost child protocol (box office)',
    'Judge names and bios (lobby)',
])

add_small_heading('What to Announce')
add_checklist(doc, [
    'Welcome & housekeeping (pre-show)',
    'Safety briefing',
    'Judge introductions',
    'Sponsor acknowledgment',
    'Category transitions',
    'Break / lunch timing',
    'End-of-session hold ("judges are deliberating")',
])

add_small_heading('Last Final Check Walk')
make_table(doc, ['Area', 'What to Verify', '\u2713'], [
    ['Backstage', 'Lights on, water stations filled, signage posted', '\u2610'],
    ['Stage', 'Clean, spike marks visible, prop storage accessible', '\u2610'],
    ['Green Room', 'Chairs, mirrors, trash cans, quiet zone sign', '\u2610'],
    ['House / Audience', 'Seats clean, lighting levels set, AC working', '\u2610'],
    ['Audio Booth', 'All tracks loaded & tested, backup files accessible', '\u2610'],
    ['Light Board', 'Cues loaded, verified with stage', '\u2610'],
    ['Judge Tables', 'Per Section 7 checklist, all ready', '\u2610'],
    ['Registration', 'Table, forms, pens, signage', '\u2610'],
    ['Medical / First Aid', 'Station visible, stocked, contact confirmed', '\u2610'],
    ['Bathrooms', 'Stocked, clean, signs clear', '\u2610'],
    ['Awards Area', '(if set up early) Trophies secure, table ready', '\u2610'],
], col_widths=[3, 9.5, 1.5])

# ═══════════════════════════════════════
# SECTION 9
# ═══════════════════════════════════════
add_page_break()
doc.add_heading('9. Preliminaries', level=2)

add_goal_logic(
    'One document that runs the day: the flow, the schedule, the production cue overview, the emcee script, breaks, lunch, and everything leading up to the judge decision.',
    'Merging schedule + cues + script into a single timeline means SM, audio, emcee, and runners all read the same page. Breaks live in the book because an unscheduled break never happens. The judge decision is the hold point.'
)

add_field('Date')

add_small_heading('Run of Show Timeline')
make_table(doc, ['Time', 'Act / Segment', 'Dur.', 'Cues', 'Emcee', 'Notes'], [
    ['', 'Doors open', '\u2014', 'House lights up', 'Welcome music', ''],
    ['', 'Top of show', '', 'Lights down', 'Welcome / housekeeping', ''],
    ['', 'Category 1 \u2014 [Name]', '', '', '', ''],
    ['', '  Act 1\u20131', '', '', 'Read lead-in', ''],
    ['', '  Act 1\u20132', '', '', '', ''],
    ['', '  ...', '', '', '', ''],
    ['', 'Break', '', 'House lights', 'Break announcement', ''],
    ['', 'Category 2 \u2014 [Name]', '', '', '', ''],
    ['', 'Lunch', '', 'Hold music', 'Lunch announcement', ''],
    ['', '...', '', '', '', ''],
    ['', 'Last act', '', '', '', ''],
    ['', 'Judges deliberate', '', 'Hold music', '"Deliberating"', 'HOLD'],
], col_widths=[2, 4, 1.2, 3, 3, 2.5])

add_small_heading('Production Cue Overview (SM + Audio + Lights)')
make_table(doc, ['Cue#', 'Moment', 'Sound', 'Light', 'SM'], [
    ['1', 'Show start', 'Fade intro \u2192 GO', 'Half \u2192 Black', 'Curtain up'],
    ['2', '[Act name]', 'Play track [file]', 'Wash [color]', '[Entry notes]'],
    ['3', 'Act exit', 'Fade out', 'Crossfade next', '[Exit notes]'],
    ['...', '', '', '', ''],
], col_widths=[1.5, 3, 4, 3.5, 3.5])

add_small_heading('Break Schedule')
make_table(doc, ['Break', 'Duration', 'Activities'], [
    ['Morning break', '__________________', '__________________'],
    ['Lunch', '__________________', '__________________'],
    ['Afternoon break', '__________________', '__________________'],
], col_widths=[3.5, 3.5, 8.5])

add_small_heading('Before Judge Decision (Hold Point)')
add_checklist(doc, [
    'All acts in the round have performed',
    'Judge sheets collected and delivered to tabulation',
    'Emcee holds the stage / hold music playing',
    'No new acts begin until tabulation confirms "ready"',
    'Results sealed / handed to emcee',
])

# ═══════════════════════════════════════
# SECTION 10
# ═══════════════════════════════════════
add_page_break()
doc.add_heading('10. Judge Review Sheets and Decision', level=2)

add_goal_logic(
    'Provide the sheets judges write on and the structure for how their decision is recorded.',
    'Sheet design determines tabulation speed. Pre-printing competitor names in act order means judges write almost nothing \u2014 fewer errors, faster counting. Every decision exists on paper, defensible later.'
)

add_small_heading('Judge Scoring Sheet Template')
p = doc.add_paragraph()
r = p.add_run('(Print one per judge, per category \u2014 pre-fill competitor names before printing.)')
r.font.size = Pt(8.5)
r.font.italic = True
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

add_mono_block(doc, '''PRELIMINARY / SEMI / FINAL  (circle one)
Category: ___________  Judge: ___________  Date: ___________

Act#  Competitor        Tech/10  Artis./10  Perf./10  Diff./10  Total/40
\u2500\u2500\u2500\u2500  \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500  \u2500\u2500\u2500\u2500\u2500\u2500  \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500  \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500  \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500  \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500
      (pre-fill first)                              \n      (pre-fill first)                              \n      (pre-fill first)                              \n      (pre-fill first)                              \n
Notes: _______________________________________________

Judge Signature: __________________''', size=8.5)

add_small_heading('Rubric Definitions')
make_table(doc, ['Criterion', 'Points', 'Description'], [
    ['Technical', '/10', 'Execution, alignment, control, precision'],
    ['Artistry', '/10', 'Musicality, expression, emotional connection'],
    ['Performance', '/10', 'Stage presence, energy, audience engagement'],
    ['Difficulty', '/10', 'Complexity of choreography, risk, innovation'],
], col_widths=[3, 2, 10.5])

add_small_heading('Tabulation Flow')
p = doc.add_paragraph()
steps = [
    'Runners collect all sheets \u2192 deliver to tabulation table (secure, quiet, no interruptions)',
    'Enter scores per competitor (drop highest/lowest if used, average remaining)',
    'Rank by total score \u2192 determine advancing number (top X per category or overall)',
    'Double-check top 3 against raw sheets',
    'Seal results in envelope \u2192 hand to emcee or stage manager',
]
for i, s in enumerate(steps, 1):
    r = p.add_run(f'{i}. {s}\n')
    r.font.size = Pt(10)

add_field('Tiebreak Rule')

# ═══════════════════════════════════════
# SECTION 11
# ═══════════════════════════════════════
add_page_break()
doc.add_heading('11. Semi Finals Preparation', level=2)

add_goal_logic(
    'Process advancing competitors, rebuild the master file, print it, and post it.',
    'A wrong name on the advancing list is public and brutal. Verify against tabulation before anything prints. The print-and-post is the official announcement \u2014 treat it as a verified act, not a chore.'
)

add_field('Semi Finals Date')

add_small_heading('Steps')
add_checklist(doc, [
    'Retrieve sealed prelims results',
    'First verification \u2014 compare advancing names against raw judge sheets',
    'Second verification \u2014 independent person cross-checks',
    'Build semi finals master file (advancing only, new act numbers)',
    'Load / verify music for advancing acts (spot-check changed files)',
    'Print competitor list (3 copies: SM, audio booth, backstage)',
    'Post at backstage bulletin board',
    'Post in lobby',
])

add_small_heading('Verification Sign-Off')
make_table(doc, ['Step', 'Completed By', 'Status'], [
    ['1st verification (vs raw sheets)', '', '\u2610'],
    ['2nd verification (independent)', '', '\u2610'],
    ['Semi master file built', '', '\u2610'],
    ['Music verified for advancing acts', '', '\u2610'],
    ['Lists printed', '', '\u2610'],
    ['Lists posted', '', '\u2610'],
], col_widths=[6.5, 5, 1.5])

# ═══════════════════════════════════════
# SECTION 12
# ═══════════════════════════════════════
add_page_break()
doc.add_heading('12. Semi Finals', level=2)

add_goal_logic(
    'One document that runs the day: flow, schedule, production cues, emcee script, breaks, up to the judge decision.',
    'This mirrors Section 9 deliberately \u2014 same structure, fewer acts, higher stakes. The crew runs on muscle memory; only the act list, script edits, and timings change. Familiar structure frees attention for what actually differs.'
)

add_field('Date')

add_small_heading('Run of Show Timeline')
make_table(doc, ['Time', 'Act / Segment', 'Dur.', 'Cues', 'Emcee', 'Notes'], [
    ['', 'Doors open', '\u2014', 'House up', 'Welcome music', ''],
    ['', 'Top of show', '', 'Lights down', 'Welcome', ''],
    ['', 'Category 1 \u2014 [Name]', '', '', '', ''],
    ['', '  ...', '', '', '', ''],
    ['', 'Lunch', '', 'Hold music', 'Announcement', ''],
    ['', '  ...', '', '', '', ''],
    ['', 'Judges deliberate', '', 'Hold music', '"Deliberating"', 'HOLD'],
], col_widths=[2, 4, 1.2, 3, 3, 2.5])

p = doc.add_paragraph()
r = p.add_run('Use the same template structure as Section 9 (Preliminaries). Only act list, category counts, and break times differ.')
r.font.italic = True
r.font.size = Pt(9.5)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

add_small_heading('Production Cue Overview')
p = doc.add_paragraph()
r = p.add_run('Same format as Section 9 \u2014 fill in semi-specific cues.')
r.font.italic = True
r.font.size = Pt(9.5)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

add_small_heading('Before Judge Decision (Hold Point)')
add_checklist(doc, [
    'All acts complete',
    'Sheets collected and delivered to tabulation',
    'Results sealed for finals prep',
])

# ═══════════════════════════════════════
# SECTION 13
# ═══════════════════════════════════════
add_page_break()
doc.add_heading('13. Finals Prep + Awards Prep', level=2)

add_goal_logic(
    'Rebuild the finals master file and stage every element of the awards ceremony before finals day starts.',
    'Finals day ends in the ceremony. Prep here covers both the competitor file and awards logistics \u2014 trophy count, envelope sealing, presenter confirmation, winner staging marks. Anything not staged here gets improvised in front of the full audience.'
)

add_field('Finals Date')

add_small_heading('Competitor Steps')
add_checklist(doc, [
    'Retrieve sealed semi finals results',
    'Verify advancing names (double-check system)',
    'Build finals master file',
    'Assign finals act numbers',
    'Print finals competitor list (3+ copies)',
    'Post at backstage + lobby',
])

add_small_heading('Awards Pre-Staging Checklist')
make_table(doc, ['Step', '\u2713'], [
    ['Count trophies/medals against advancing categories', '\u2610'],
    ['Organize awards by category in presentation order', '\u2610'],
    ['Stage awards table near stage wing', '\u2610'],
    ['Prepare sealed envelopes per category', '\u2610'],
    ['Confirm award presenters + name pronunciation', '\u2610'],
    ['Stage marks for winners (tape X: 1st center, 2nd right, 3rd left)', '\u2610'],
    ['Photo backdrop / area marked', '\u2610'],
], col_widths=[14.5, 1.5])

add_small_heading('Finals Stage Diagram (with Award Marks)')
add_mono_block(doc, '''        [Upstage]
        3rd(X)   2nd(X)   1st(X)
                  |
           [Awards Table]

        [Downstage]''', size=9.5)

add_small_heading('Awards Inventory')
make_table(doc, ['Category', '1st', '2nd', '3rd', 'Qty Needed', 'Qty Available'], [
    ['', '', '', '', '', ''],
    ['', '', '', '', '', ''],
    ['', '', '', '', '', ''],
    ['', '', '', '', '', ''],
], col_widths=[3, 1.5, 1.5, 1.5, 2.5, 2.5])

add_small_heading('Awards Presentation Order')
make_table(doc, ['Order', 'Category', 'Presenter', 'Notes'], [
    ['1', '', '', ''],
    ['2', '', '', ''],
    ['3', '', '', ''],
    ['4', '', '', ''],
    ['5', '', '', ''],
], col_widths=[1.5, 4.5, 4.5, 5])

# ═══════════════════════════════════════
# SECTION 14
# ═══════════════════════════════════════
add_page_break()
doc.add_heading('14. Finals', level=2)

add_goal_logic(
    'One document that runs the day: flow, schedule, production cues, emcee script, breaks, rewards ceremony spacing, lunch, up to the judge decision.',
    'Same skeleton as prelims and semis, plus "rewards ceremony spacing" \u2014 the stage must physically transform for the ceremony (winner marks, trophy table, photo sightlines), and that transformation is planned inside the day\'s timeline, not squeezed in.'
)

add_field('Date')

add_small_heading('Run of Show Timeline')
make_table(doc, ['Time', 'Segment', 'Dur.', 'Cues', 'Emcee', 'Notes'], [
    ['', 'Doors open', '\u2014', 'House up', 'Welcome music', ''],
    ['', 'Top of show', '', 'Lights down', 'Welcome', ''],
    ['', 'Finals performance block', '', '', '', ''],
    ['', '  ...', '', '', '', ''],
    ['', 'Finals end', '', 'Blackout', '"Deliberating"', ''],
    ['', 'Stage transform for awards', '', 'Work lights', 'Hold announcement', 'Move table + marks'],
    ['', 'Awards ceremony', '', '', 'Per Section 15', ''],
], col_widths=[2, 4, 1.2, 3, 3, 2.5])

add_small_heading('Rewards Ceremony Spacing')
make_table(doc, ['Action', 'Dur.', 'Who', 'Cue'], [
    ['Last act finishes', '\u2014', '\u2014', 'Blackout'],
    ['Work lights up', '30s', 'Lights', 'SM call'],
    ['Awards table moved to stage', '2 min', 'Crew', ''],
    ['Podium marks taped', '1 min', 'Crew', ''],
    ['Trophies arranged on table', '2 min', 'Awards lead', ''],
    ['Emcee takes position', '30s', 'Emcee', ''],
    ['House lights down', '\u2014', 'Lights', 'SM call'],
    ['Awards ceremony begins', '\u2014', 'Emcee', 'Per Section 15'],
], col_widths=[5.5, 1.5, 2.5, 2.5])

add_small_heading('Before Judge Decision (Final Hold)')
add_checklist(doc, [
    'All finals acts complete',
    'Judge sheets collected',
    'Tabulation underway',
    'Stage transformation can run parallel to tabulation (if separate crew)',
    'If same crew \u2014 wait for results before transforming',
])

# ═══════════════════════════════════════
# SECTION 15
# ═══════════════════════════════════════
add_page_break()
doc.add_heading('15. Awards', level=2)

add_goal_logic(
    'Run the ceremony: exact order of presentation, who announces, winner flow on and off stage, photo moments.',
    'The ceremony is the audience\'s last memory. A scripted, tightly-run ceremony beats a rambling one. Ending clean and on time sends everyone home remembering the show worked.'
)

add_small_heading('Ceremony Run of Show')
make_table(doc, ['Step', 'Action', 'Script / Cue', 'Dur.'], [
    ['1', 'House lights down', '', '5s'],
    ['2', 'Emcee welcome \u2192 awards transition', '[Script]', '30s'],
    ['3', 'Presenter intro', '"Please welcome..."', '10s'],
    ['4', 'Category 1 \u2014 3rd place', 'Envelope opened, name read', '20s'],
    ['', 'Winner crosses stage, takes mark', 'Photo', '15s'],
    ['', '2nd place', 'Same flow', '20s'],
    ['', '1st place', 'Same flow', '30s'],
    ['', 'Category photo (all three)', '', '20s'],
    ['5', 'Next category', 'Repeat step 4', '\u2014'],
    ['...', '...', '...', '...'],
    ['N', 'Final category / Grand Champion', '', ''],
    ['N+1', 'Final sponsor acknowledgment', '', '30s'],
    ['N+2', 'Closing remarks', '', '30s'],
    ['N+3', 'House lights up', '', '5s'],
], col_widths=[1.5, 5, 5.5, 1.5])

add_small_heading('Standardized Winner Flow')
add_mono_block(doc, '''Name called \u2192 Walk from wing to center mark \u2192
Turn to audience, accept award \u2192 Photographer captures \u2192
Remain on mark until all three placed \u2192
Category photo \u2192 Exit together''', size=9)

add_small_heading('Photo Moments')
add_checklist(doc, [
    'Individual winner with award',
    'Category top 3 together',
    'Presenter + winners (if desired)',
    'Full cast / studio photo (after ceremony or during exit)',
])

add_field('Ceremony Duration Target')

# ═══════════════════════════════════════
# POST-EVENT NOTES
# ═══════════════════════════════════════
add_page_break()
doc.add_heading('Post-Event Notes', level=2)

p = doc.add_paragraph()
r = p.add_run('Complete after the event and file this book with notes clipped inside the front cover. The book improves every season.')
r.font.italic = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

add_small_heading('What ran late?')
p = doc.add_paragraph()
r = p.add_run('_' * 70)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

add_small_heading('What broke / failed?')
p = doc.add_paragraph()
r = p.add_run('_' * 70)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

add_small_heading('What confused volunteers?')
p = doc.add_paragraph()
r = p.add_run('_' * 70)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

add_small_heading('What to change next time?')
p = doc.add_paragraph()
r = p.add_run('_' * 70)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

add_small_heading('New sections to add / sections to remove?')
p = doc.add_paragraph()
r = p.add_run('_' * 70)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── Save ──
output_path = '/Users/patricktrang/opencode-2607/dance-competition-promptbook.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
